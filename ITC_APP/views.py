from django.shortcuts import render,redirect, get_object_or_404
from .models import Anggota,Sertifikat,Bidang,Pengurus,Program,Event,Notification
from .forms import AnggotaForm,SertifikatForm,ProgramForm,PengurusForm,BidangForm
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.core.paginator import Paginator
from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.utils import timezone
import json
from django.http import JsonResponse
from datetime import date
from django.http import HttpResponse
from openpyxl import Workbook
import pywhatkit as kit
import datetime
from django.db.models import Count
from django.utils import timezone
from django.db.models.functions import ExtractWeekDay
from docx import Document
from docx.shared import Inches


def notif(request):
    bid_nav = Bidang.objects.all()
    notif = Notification.objects.all().order_by('-id')
    context ={
        'bid_nav':bid_nav,
        'notif':notif,
    }
    return render(request,'notif.html',context)

@login_required()
@csrf_exempt
def notif_remove(request,id):
    notif = get_object_or_404(Notification, id=id)
    notif.delete()
    return redirect('dashboard:notif')

@login_required()
@csrf_exempt
def dashboard(request):
    jlh_anggota = Anggota.objects.all().count()
    jlh_verifikasi = Anggota.objects.filter(diverifikasi=True).count()
    bid_nav = Bidang.objects.all()
    program = Event.objects.all().count()
    notif = Notification.objects.all().count()

    data_pengurus_per_bidang = []
    bidangs = Bidang.objects.all()

    # card jumlah masing-masing pengrusu per bidang
    for bidang in bidangs:
        jumlah_pengurus = Pengurus.objects.filter(bidang=bidang).count()
        data_pengurus_per_bidang.append({
            'nama_bidang': bidang.nama_bidang,
            'jumlah_pengurus': jumlah_pengurus,
        })

    # notifikasi
    days_of_week = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']

    # Menghitung jumlah notifikasi per hari
    notifikasi_per_hari = Notification.objects.annotate(day_of_week=ExtractWeekDay('timestamp')).values('day_of_week').annotate(total=Count('id'))

    # Menyusun hasil query ke dalam format yang sesuai untuk grafik
    data_notifikasi = [0] * 7
    for item in notifikasi_per_hari:
        day_index = item['day_of_week'] - 1  # Mengonversi indeks hari ke format Python (0-6)
        data_notifikasi[day_index] = item['total']


    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()
    context = {
        'bid_nav': bid_nav,
        'jlh_anggota': jlh_anggota,
        'jlh_verifikasi': jlh_verifikasi,
        'data_pengurus_per_bidang': data_pengurus_per_bidang,
        'notifications': notifications,
        'notif': notif,
        'program': program,
        'notif': notif,
        'days_of_week': days_of_week,
        'data_notifikasi': data_notifikasi
    }

    return render(request, 'dashboard.html', context)

# ANGGOTA
@login_required()
@csrf_exempt
def anggota_create(request):
    bid_nav = Bidang.objects.all()
    if request.method == 'POST':
        form = AnggotaForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, " Data Berhasil di Tambahkan!")

            # Create notification
            Notification.objects.create(
                message=f'Anggota baru "{form.instance.nama}" ditambahkan!',
                notification_type='new_member'
            )

            return redirect('dashboard:anggota')
    else:
        form = AnggotaForm()
    
    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context = {
        'form': form,
        'bid_nav': bid_nav,
        'notifications': notifications,
        'notif': notif,
    }
    return render(request, 'anggota_crate.html', context)

@login_required()
@csrf_exempt
def anggota(request):
    data = Anggota.objects.all().order_by('-id')
    bid_nav = Bidang.objects.all()

    #search 
    search_query = request.GET.get('cari')
    if search_query:
        data = data.filter(
            Q(nama__icontains=search_query) |
            Q(nim__icontains=search_query) |
            Q(alamat__icontains=search_query) |
            Q(diverifikasi__icontains=search_query)
        )

    # pagination
    paginator = Paginator(data, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context = {
        'datas': page_obj,
        'bid_nav': bid_nav,
        'search_query': search_query,
        'notifications': notifications,
        'notif': notif,
    }

    if request.method == 'POST':
        anggota_id = request.POST.get('anggota_id')
        anggota = Anggota.objects.get(id=anggota_id)
        anggota.diverifikasi = not anggota.diverifikasi
        anggota.save()

    return render(request, 'anggota.html', context)

@login_required()
@csrf_exempt
def anggota_update(request,id):
    bid_nav = Bidang.objects.all()
    anggota = get_object_or_404(Anggota, id=id)
    if request.method == 'POST':
        form = AnggotaForm(request.POST, instance=anggota)

        if form.is_valid():
            form.save()
            messages.success(request, " Data Berhasil di Update!")
            return redirect('dashboard:anggota')
    else:
        form = AnggotaForm(instance=anggota)

    
    context ={
        'form': form,
        'bid_nav': bid_nav,
    }

    return render(request, 'anggota_crate.html',context )

@login_required()
@csrf_exempt
def anggota_detail(request,id):
    bid_nav = Bidang.objects.all()
    anggota = get_object_or_404(Anggota, id=id)

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()
    context = {
        'anggota': anggota,
        'bid_nav': bid_nav,
        'notifications': notifications,
        'notif': notif,
    }
    return render(request,'anggota_detail.html',context)

@login_required()
@csrf_exempt
def anggota_delete(request,id):
    anggota = get_object_or_404(Anggota, id=id)
    anggota.delete()
    return redirect('dashboard:anggota')

#bidang
@login_required()
def bidang_list(request):
    bid_nav = Bidang.objects.all()

   
    context = {
        'bid_nav': bid_nav,
      
        }
    return render(request, 'snippets/navbar.html',context )

@login_required()
def bidang(request):
    bid_nav = Bidang.objects.all()
    search_query = request.GET.get('cari')

    if search_query:
        bid_nav = Bidang.objects.filter(nama_bidang__icontains=search_query)
    else:
        bid_nav = Bidang.objects.all()

    paginator = Paginator(bid_nav, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context = {
        'bidang': page_obj,
        'bid_nav':bid_nav,
        'notifications':notifications,
        'notif':notif,
    }
    return render(request, 'bidang.html', context)

@login_required()
@csrf_exempt
def bidang_create(request):
    bid_nav = Bidang.objects.all()
    if request.method == 'POST':
        form = BidangForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, " Data Berhasil di Tambah!")
            return redirect('dashboard:bidang_list')
    else:
        form = BidangForm()

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context = {
        'form': form,
        'bid_nav': bid_nav,
        'notifications': notifications,
        'notif': notif,
    }
    return render(request, 'bidang_create.html',context)

@login_required()
@csrf_exempt
def bidang_update(request,id):
    bid_nav = Bidang.objects.all()
    bidangs = get_object_or_404(Bidang, id=id)
    if request.method == 'POST':
        form = BidangForm(request.POST, instance=bidangs)

        if form.is_valid():
            form.save()
            messages.success(request, " Data Berhasil di Update!")
            return redirect('dashboard:bidang_list')
    else:
        form = BidangForm(instance=bidangs)
    
    context ={
        'form': form,
        'bid_nav': bid_nav,
    }

    return render(request, 'bidang_create.html',context )

@login_required()
@csrf_exempt
def bidang_delete(request,id):
    bidang = get_object_or_404(Bidang, id=id)
    bidang.delete()
    return redirect('dashboard:bidang_list')

# pengurus
@login_required()  
def pengurus_list(request, id):
    bid_nav = Bidang.objects.all()
    bidangs = get_object_or_404(Bidang, id=id)
    pengurus = bidangs.bidang.all()

    query = request.GET.get('cari')
    if query:
        pengurus = pengurus.filter(
            Q(nama_pengurus__icontains=query) | Q(status__icontains=query)
        )

    paginator = Paginator(pengurus,5)  
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context = {
        'bidangs': bidangs, 
        'bid_nav': bid_nav, 
        'pengurus': page_obj, 
        'query': query,
        'notifications': notifications,
        'notif': notif,
    }

    return render(request, 'pengurus.html', context)

@login_required()
@csrf_exempt
def pengurus_create(request):
    bid_nav = Bidang.objects.all()
    if request.method == 'POST':
        form = PengurusForm(request.POST)
        if form.is_valid():
            pengurus = form.save()
            messages.success(request, "Data Berhasil di Tambah!")
            return redirect('dashboard:pengurus', id=pengurus.bidang.id)
    else:
        form = PengurusForm()
    
    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context = {
        'form': form,
        'bid_nav': bid_nav,
        'notifications': notifications,
        'notif': notif,
    }
    return render(request, 'pengurus_crate.html',context)

@login_required()
@csrf_exempt
def pengurus_update(request,id):
    bid_nav = Bidang.objects.all()
    pengurus = get_object_or_404(Pengurus, id=id)
    if request.method == 'POST':
        form = PengurusForm(request.POST, instance=pengurus)

        if form.is_valid():
            form.save()
            messages.success(request, " Data Berhasil di Update!")
            return redirect('dashboard:anggota')
    else:
        form = PengurusForm(instance=pengurus)
    
    context ={
        'form': form,
        'bid_nav': bid_nav,
    }

    return render(request, 'pengurus_crate.html',context )

@login_required()
@csrf_exempt
def pengurus_delete(request,id):
    pengurus = get_object_or_404(Pengurus, id=id)
    pengurus.delete()
    return redirect('dashboard:anggota')

# program
@login_required()
@csrf_exempt
def calendar_view(request):
    events = Event.objects.all()
    bid_nav = Bidang.objects.all()

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()
    
    context =  {
        'bid_nav':bid_nav,
        'events': events,
        'event': events,
        'notifications': notifications,
        'notif': notif,
    }
    return render(request, 'program_calendar.html',context)

@login_required()
@csrf_exempt
def all_events(request):                                                                                                 
    all_events = Event.objects.all()                                                                                    
    out = []                                                                                                             
    for event in all_events:                                                                                             
        out.append({ 
            'id': event.id,                                                                                                    
            'title': event.title,                                                                                                                                                                                       
            'start': event.start_date.strftime("%m/%d/%Y, %H:%M:%S"),                                                         
            'end': event.end_date.strftime("%m/%d/%Y, %H:%M:%S"),                                                         
        })                                                                                                               
                                                                                                                      
    return JsonResponse(out, safe=False) 

@login_required()
@csrf_exempt
def add_event(request):
    if request.method == 'POST':
        title = request.POST['title']
        start_date = request.POST['start_date']
        end_date = request.POST['end_date']

        event = Event(title=title, start_date=start_date, end_date=end_date)
        event.save()

        # Create notification
        Notification.objects.create(
            message=f'Agenda "{title}" ditambahkan!',
            notification_type='event_added'
        )

        return redirect('dashboard:program_calendar')

@login_required()
@csrf_exempt
def remove_event(request, event_id):
    try:
        event = Event.objects.get(id=event_id)
        event.delete()

        data = {'success': True}
    except Event.DoesNotExist:
        data = {'success': False, 'message': 'Acara dengan ID yang diberikan tidak ada.'}

    return JsonResponse(data)

@login_required()
@csrf_exempt
def update_program(request):
    id = request.GET.get("id", None)
    title = request.GET.get("title", None)
    start = request.GET.get("start", None)
    end = request.GET.get("end", None)

    try:
        event = Event.objects.get(id=id)
        event.title = title
        event.start_date = start
        event.end_date = end
        event.save()

        data = {'success': True}
    except Event.DoesNotExist:
        data = {'success': False, 'message': 'Acara dengan ID yang diberikan tidak ada.'}

    return JsonResponse(data)

# laporan
@login_required()
@csrf_exempt
def laporan(request):
    bid_nav = Bidang.objects.all()
    anggota = Anggota.objects.all().order_by('-diverifikasi')

    today = timezone.now().date()

    # Menghitung jumlah anggota yang diverifikasi per hari
    diverifikasi_per_hari = Anggota.objects.filter(diverifikasi=True, tanggal_diverifikasi__gte=today).values('tanggal_diverifikasi').annotate(total=Count('id'))

    # Menghitung jumlah anggota baru per hari
    anggota_baru_per_hari = Notification.objects.filter(notification_type='new_member', timestamp__date__gte=today).values('timestamp__date').annotate(total=Count('id'))

    # Mengonversi hasil query ke format yang sesuai untuk grafik
    data_diverifikasi = [{'tanggal': item['tanggal_diverifikasi'].strftime("%d %B %Y"), 'jumlah': item['total']} for item in diverifikasi_per_hari if item['tanggal_diverifikasi']]
    data_anggota_baru = [{'tanggal': item['timestamp__date'].strftime("%d %B %Y"), 'jumlah': item['total']} for item in anggota_baru_per_hari if item['timestamp__date']]

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context={
        'bid_nav':bid_nav,
        'anggota':anggota,
        'data_diverifikasi': data_diverifikasi,
        'data_anggota_baru': data_anggota_baru,
        'notifications': notifications,
        'notif': notif,
    }
    return render(request,'laporan.html',context)

# chat
@csrf_exempt
def chat(request):
    bid_nav = Bidang.objects.all()
    anggota_list = Anggota.objects.all().order_by('diverifikasi')

    if request.method == 'POST':
        nomor_list = request.POST.getlist('nomor')
        pesan = request.POST.get('pesan')

        now = datetime.datetime.now()
        jam = now.hour
        menit = now.minute + 1  # waktu pengiriman
        for nomor in nomor_list:
            kit.sendwhatmsg(nomor, pesan, jam, menit)  

    notifications = Notification.objects.all().order_by('-id')[:5]
    notif = Notification.objects.all().count()

    context={
        'bid_nav':bid_nav,
        'anggota_list':anggota_list,
        'notifications':notifications,
        'notif':notif,
    }
    return render(request,'form_chat.html',context)

# sertifikat
def sertifikat_upload(request):
    if request.method == 'POST':
        form = SertifikatForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('lihat_sertifikat')
    else:
        form = SertifikatForm()
    
    return render(request, 'sertifikat_upload.html', {'form': form})

def lihat_sertifikat(request):
    sertifikat_list = Sertifikat.objects.all()
    return render(request, 'sertifikat.html', {'data': sertifikat_list})

def export_excel(request):
    anggota = Anggota.objects.all().order_by('-diverifikasi')

    workbook = Workbook()
    worksheet = workbook.active

    # Set header row
    worksheet.append(['NO', 'NAMA', 'J/K', 'ALAMAT', 'QR','PRODI', 'SEMESTER', 'STATUS'])

    # Populate data
    for index, data in enumerate(anggota, start=1):
        jalur_gambar_qr = data.qr.path  # Diasumsikan data.qr adalah ImageField
        tautan_gambar_qr = f'<img src="{jalur_gambar_qr}" alt="QR Code">'
        worksheet.append([index, data.nama, data.jk, data.alamat, tautan_gambar_qr, data.prodi, data.semester, "Diverifikasi" if data.diverifikasi else "Verifikasi"])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=laporan.xlsx'
    workbook.save(response)

    return response

def update_cetak_status(request):
    if request.method == 'POST':
        anggota_id = request.POST.get('anggota_id')
        is_checked = request.POST.get('isChecked')

        anggota = Anggota.objects.get(id=anggota_id)
        anggota.cetak = is_checked
        anggota.save()

        return JsonResponse({'message': 'Status cetak diperbarui'}, status=200)
    else:
        return JsonResponse({'message': 'Metode permintaan tidak valid'}, status=400)

def print_preview(request):
    if request.method == 'POST':
        selected_ids = json.loads(request.POST.get('selected_ids', '[]'))
        Anggota.objects.filter(id__in=selected_ids).update(cetak=True)

        return HttpResponse(status=200)

    return HttpResponseBadRequest('Invalid request method')

def data_anggota(request):
    bid_nav = Bidang.objects.all()
    anggota = Anggota.objects.all().order_by('-diverifikasi')

    context ={
        'bid_nav':bid_nav,
        'anggota':anggota
    }

    return render(request,'data_anggota.html',context)

def export_to_word(request):
    anggota = Anggota.objects.all().order_by('-diverifikasi')

    # Inisialisasi dokumen Word
    doc = Document()

    # Tambahkan judul
    doc.add_heading('Daftar Anggota', level=1)

    # Tambahkan tabel
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.autofit = False

    # Atur lebar kolom
    widths = (15, 30, 10, 30, 20, 10, 10, 20)  # Sesuaikan lebar kolom sesuai kebutuhan
    for i, width in enumerate(widths):
        table.columns[i].width = width

    # Tambahkan header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'NO'
    header_cells[1].text = 'NAMA'
    header_cells[2].text = 'J/K'
    header_cells[3].text = 'ALAMAT'
    header_cells[4].text = 'PRODI'
    header_cells[5].text = 'SM'
    header_cells[6].text = 'QR'
    header_cells[7].text = 'STATUS'

    # Populate data
    for index, data in enumerate(anggota, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = data.nama
        row_cells[2].text = data.jk
        row_cells[3].text = data.alamat
        row_cells[4].text = data.prodi
        row_cells[5].text = str(data.semester)
        # Sisipkan gambar
        qr_img_path = data.qr.path
        row_cells[6].paragraphs[0].add_run().add_picture(qr_img_path, width=Inches(1.25))
        # Tambahkan status
        status = "Diverifikasi" if data.diverifikasi else "Verifikasi"
        row_cells[7].text = status

    # Simpan dokumen
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=daftar_anggota.docx'
    doc.save(response)
    return response
